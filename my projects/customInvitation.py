#! python3
#  a program that would generate a Word document with custom invitations
import docx

def customInv():
    doc = docx.Document('invitationLetterCustom.docx')
    file = open('guests.txt')
    fileList = file.readlines()
    for line in range(len(doc.paragraphs)):
        doc.paragraph[-1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
        doc.paragraph[line] = 
    for a in range(len(fileList)):
        doc.paragraphs[2].add_run(fileList[a])

customInv()