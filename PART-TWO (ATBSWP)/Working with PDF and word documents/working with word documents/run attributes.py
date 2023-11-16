import docx

doc = docx.Document('C:\\Users\\user\\Desktop\\PYTHON\\demo.docx')
doc.paragraphs[0].text
'Document Title'
doc.paragraphs[0].style
'''_ParagraphStyle('Title') id: 2221900876288'''
doc.paragraphs[0].style = 'Normal'
doc.paragraphs[1].text
'A plain paragraph with some bold and some italic'
(doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text)
('A plain paragraph with', ' some ', 'bold', ' and some ')
doc.paragraphs[1].runs[0].style = 'QuoteChar'

'''Warning (from warnings module):
  File "C:\Users\user\AppData\Local\Programs\Python\Python310\lib\site-packages\docx\styles\styles.py", line 139
    return self._get_style_id_from_style(self[style_name], style_type)
UserWarning: style lookup by style_id is deprecated. Use style name as key instead.'''
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save('C:\\Users\\user\\Desktop\\PYTHON\\restyled.docx')
