import docx

doc = docx.Document('C:\\Users\\user\\Desktop\\PYTHON\\demo.docx')
len(doc.paragraphs)
7
doc.paragraphs[0].text
'Document Title'
doc.paragraphs[1].text
'A plain paragraph with some bold and some italic'
len(doc.paragraphs[1].runs)
5
doc.paragraphs[1].runs[0].text
'A plain paragraph with'
doc.paragraphs[1].runs[1].text
' some '
doc.paragraphs[1].runs[2].text
'bold'
doc.paragraphs[1].runs[3].text
' and some '
