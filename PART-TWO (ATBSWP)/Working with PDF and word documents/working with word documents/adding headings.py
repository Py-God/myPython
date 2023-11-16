import docx

doc = docx.Document()
doc.add_heading('Header 0', 0)
'''<docx.text.paragraph.Paragraph object at 0x000001B4F133C820>'''
doc.add_heading('Header 1', 1)
'''<docx.text.paragraph.Paragraph object at 0x000001B4F133CA30>'''
doc.add_heading('Header 2', 2)
'''<docx.text.paragraph.Paragraph object at 0x000001B4F133C850>'''
doc.add_heading('Header 3', 3)
'''<docx.text.paragraph.Paragraph object at 0x000001B4F133C910>'''
doc.add_heading('Header 4', 4)
'''<docx.text.paragraph.Paragraph object at 0x000001B4F133CA30>'''
doc.save('C:\\Users\\user\\Desktop\\PYTHON\\headings.docx')
