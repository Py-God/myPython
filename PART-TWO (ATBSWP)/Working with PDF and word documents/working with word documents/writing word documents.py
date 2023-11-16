import docx

doc = docx.Document()
doc.add_paragraph('Hello world!')
'''<docx.text.paragraph.Paragraph object at 0x00000249385EC9A0>'''
doc.save('helloworld.docx')
doc = docx.Document('helloworld.docx')
paraObj1 = doc.add_paragraph('This is a second paragraph.')

paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph.')
'''<docx.text.run.Run object at 0x00000249385EE230>'''
doc.save('multipleParagraphs.docx')
doc = docx.Document('multipleParagraphs.docx')
doc.add_paragraph('Hello world!', 'Title')  # adding a style from add_paragrapgh
'''<docx.text.paragraph.Paragraph object at 0x00000249385ECF40>'''
doc.save('Helloworld.docx')
