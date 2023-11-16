import docx

doc = docx.Document()
doc.add_paragraph('This is on the first page!')
'''<docx.text.paragraph.Paragraph object at 0x0000027F891309A0>'''
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')
'''<docx.text.paragraph.Paragraph object at 0x0000027F89130B50>'''
doc.save('C:\\Users\\user\\Desktop\\PYTHON\\twoPage.docx')
